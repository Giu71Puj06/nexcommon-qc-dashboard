# -*- coding: utf-8 -*-
"""
Generatore Schede Ispettive (UNICO) - v8.6.7
--------------------------------------------------------------------------

Cosa fa (in un solo script):
1) Genera le schede ispettive per DISCIPLINA a partire da:
   - ELENCO ELABORATI (.xlsx)
   - ToDo Trimble export (.xlsx)
   - Report coerenze/elaborati (.xlsx)  -> per compilare "TITOLO ELABORATO"
   - Template Word (.docx)

2) Gestione immagini:
   - Recupera immagini dal ToDo se presenti colonne Foto/Immagini/Snapshot (path o URL)
   - E/O da una cartella foto selezionata in GUI (root), con sottocartelle disciplina:
       <FotoRoot>\<DISCIPLINA>\IT22-65_Foto1.png ...
   - Le immagini vengono salvate/rinominate nell'output (per disciplina) con regola:
       LABEL_Foto.png (una sola) oppure LABEL_Foto1.png, LABEL_Foto2.png...

3) Inserimento immagini nel DOCX (come richiesto):
   - In tabella (prima colonna) resta SOLO "NC-LABEL" / "OSS-LABEL".
   - Se quella NC/OSS ha immagini -> diventa un LINK INTERNO che rimanda all'immagine in appendice.
   - A fine documento viene creata la sezione "ALLEGATI IMMAGINI":
       - immagini CENTRATE
       - UNA SOLA didascalia sotto ogni immagine (centrata): "NC-IT22-65 - <descrizione>"

4) Regole:
   - Se "NESSUN RILIEVO" nel ToDo -> NON inserisce la riga.
   - Codice rilievo:
       Tag OSS -> OSS-LABEL
       altrimenti -> NC-LABEL

NOVITA v8.6.7:
- Nella tabella finale, la colonna "ASSENZA NC/OSS" viene valorizzata con "X" SOLO se il documento non ha alcuna NC e alcuna OSS.

- Colonna "ISPETTORE" nella tabella rilievi ora mostra SOLO LE INIZIALI
  (es. "Stefano Arcangelelli" -> "SA", "Ing. O. Bellaroba" -> "OB").
  Il campo "Nome Redattore" in prima pagina resta invariato (nome esteso).
- Nella maschera GUI sono stati aggiunti due campi opzionali:
    * "Data ricezione elaborati" (formato gg/mm/aaaa o gg.mm.aaaa)
    * "Data emissione scheda ispettiva" (formato gg/mm/aaaa o gg.mm.aaaa)
  Se compilati sovrascrivono rispettivamente:
    - il valore letto da ELENCO_ELABORATI per la prima pagina
    - il testo della data nella tabella revisioni (Tabella 0)
- Titoli ToDo con piu codici (es. "RMS_G_ESE_COR_ARC_030_00; RMS_G_ESE_COR_ARC_032_00")
  ora vengono espansi: il flag NC/OSS/NESSUN viene applicato a ogni singolo codice
  nel riepilogo finale.
- Titoli ToDo che contengono un nome cartella (es. "RE00_E- ARCHITETTURA STATO DI FATTO")
  con tag "Nessun rilievo": il flag NESSUN viene applicato a tutti i documenti
  della disciplina non ancora flaggati (presi da ELENCO/Report). La voce cartella
  NON compare piu nel riepilogo (allinea il conteggio).

Note:
- Richiede: pip install pandas openpyxl python-docx docxtpl requests (requests solo se usi URL immagini)
"""

import os
import sys
import re
import unicodedata
import pathlib
import shutil
import pandas as pd
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox


def resource_path(relative_path: str) -> str:
    """Percorso risorsa compatibile con PyInstaller (onefile/onedir).

    - In sviluppo: usa la cartella dello script.
    - In EXE PyInstaller: usa la cartella temporanea _MEIPASS.
    """
    base_path = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional (solo se nel ToDo ci sono URL)
try:
    import requests
except Exception:
    requests = None

IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".bmp")


# ============================================================
# UTILITY
# ============================================================
def safe_str(v):
    try:
        return "" if pd.isna(v) or v is None else str(v).strip()
    except Exception:
        return str(v).strip()


def safe_date_str(v, fmt="%d/%m/%Y"):
    """Convert Excel/Pandas dates to a clean date string (no time)."""
    try:
        if v is None or (hasattr(pd, "isna") and pd.isna(v)):
            return ""
        # pandas Timestamp / datetime
        if isinstance(v, (pd.Timestamp, datetime.datetime)):
            return v.date().strftime(fmt)
        if isinstance(v, datetime.date):
            return v.strftime(fmt)
        # strings like '2025-11-12 00:00:00' or '2025-11-12'
        s = str(v).strip()
        for f in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                d = datetime.datetime.strptime(s, f)
                return d.date().strftime(fmt)
            except Exception:
                pass
        return s
    except Exception:
        return safe_str(v)


def tr_from_label(label: str) -> str:
    """Return 'TR-<numero>' extracted from Trimble ToDo Label.

    Examples:
      'OSS-IT22-98' -> 'TR-98'
      'TR-011'      -> 'TR-011'
      'NC-XYZ-7'    -> 'TR-7'
    """
    s = safe_str(label)
    if not s:
        return ""
    # Prefer trailing digits (keeps leading zeros)
    m = re.search(r"(\d+)\s*$", s)
    if m:
        num = m.group(1)
        return f"TR-{num}"
    # Fallback: scan tokens from the end
    parts = re.split(r"[-_]", s)
    for p in reversed(parts):
        if p.isdigit():
            return f"TR-{p}"
    return s


def normalize_text(s: str) -> str:
    return re.sub(r"\s+", " ", safe_str(s).upper()).strip()


def split_assignees(value: str) -> list[str]:
    """Split the Assignee(s) field into multiple discipline names.

    Trimble exports can contain multiple disciplines separated by ';' or ',' (and sometimes newlines or '|').
    Example: "Documenti generali; Ambiente e vincoli" -> ["Documenti generali", "Ambiente e vincoli"]
    """
    s = safe_str(value)
    if not s:
        return []
    # normalize separators to semicolon
    parts = re.split(r"[;,\n\r\t\|]+", s)
    out = []
    for p in parts:
        p2 = re.sub(r"\s+", " ", p).strip()
        if p2:
            out.append(p2)
    # de-dup while preserving order
    seen=set()
    uniq=[]
    for p in out:
        k=normalize_text(p)
        if k in seen:
            continue
        seen.add(k)
        uniq.append(p)
    return uniq


def primary_redattore(value: str) -> str:
    """Return the first (primary) redattore name from ELENCO ELABORATI cell value.
    Accepts separators ';' and newlines.
    """
    s = safe_str(value)
    if not s:
        return ""
    parts = [p.strip() for p in re.split(r"[;\n\r]+", s) if p.strip()]
    return parts[0] if parts else ""


def infer_report_cols(df: pd.DataFrame) -> dict:
    """Try to infer key columns from report coerenze/non-coerenze."""
    cols = list(df.columns)

    def pick(preds):
        for c in cols:
            lc = str(c).lower()
            if any(p in lc for p in preds):
                return c
        return None

    col_codice = pick(["codice", "code"])
    col_titolo = pick(["titolo", "title"])
    col_rev = pick(["revis", "rev"])
    col_disc = pick(["disciplina", "assignee"])
    # more specific first
    col_sp = pick(["codice_sp", "codicesp"])
    if col_sp is None:
        # sometimes just 'sp'
        col_sp = pick([" sp", "_sp", "sp_"]) or pick(["sp"])
    return {
        "codice": col_codice,
        "titolo": col_titolo,
        "rev": col_rev,
        "disciplina": col_disc,
        "codice_sp": col_sp,
    }


def docs_from_report_for_disciplina(report_df: pd.DataFrame, disciplina: str, codice_sp_disc: str) -> list[dict]:
    """Return list of docs (code, rev, titolo) that belong to the given disciplina.

    Strategy:
    1) If report has 'disciplina' column -> filter by disciplina (case/space insensitive, containment allowed).
    2) Else if report has Codice_SP/SP column -> filter by codice_sp_disc.
    3) Else -> return [] (caller will fallback).
    """
    if report_df is None or report_df.empty:
        return []

    cols = infer_report_cols(report_df)
    col_cod = cols["codice"]
    if not col_cod:
        return []

    col_tit = cols["titolo"]
    col_rev = cols["rev"]
    col_disc = cols["disciplina"]
    col_sp = cols["codice_sp"]

    df = report_df.copy()

    # Filter
    if col_disc:
        kd = normalize_text(disciplina)
        kd_comp = re.sub(r"\s+", "", kd)
        def _match_disc(v):
            kv = normalize_text(v)
            kv_comp = re.sub(r"\s+", "", kv)
            return bool(kd_comp) and (kd_comp in kv_comp or kv_comp in kd_comp)
        mask = df[col_disc].apply(_match_disc)
        df = df[mask]
    elif col_sp and codice_sp_disc:
        ksp = normalize_text(codice_sp_disc)
        def _match_sp(v):
            return normalize_text(v) == ksp
        df = df[df[col_sp].apply(_match_sp)]
    else:
        return []

    out = []
    for _, rr in df.iterrows():
        code = safe_str(rr.get(col_cod, ""))
        if not code:
            continue
        out.append({
            "code": code,
            "rev": safe_str(rr.get(col_rev, "")) if col_rev else "",
            "titolo": safe_str(rr.get(col_tit, "")) if col_tit else "",
        })

    # de-dup by normalized code, keep first
    seen=set()
    uniq=[]
    for d in out:
        nk = normalize_code(d["code"])
        if not nk or nk in seen:
            continue
        seen.add(nk)
        uniq.append(d)
    return uniq

def normalize_code(s: str) -> str:
    s = unicodedata.normalize("NFKD", safe_str(s))
    s = s.upper().replace("-", "_").replace(" ", "_")
    s = re.sub(r"[^\w]", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s

def strip_revision_suffix(code: str) -> str:
    # rimuove _00, -00, _01, -01, ecc (fine stringa)
    return re.sub(r"([_-]\d{2})$", "", safe_str(code))


def format_revision_value(rev: str, code: str = "") -> str:
    """Normalizza la revisione nel formato 'Rev_00'.

    Priorita:
    1) usa il valore REV passato, se presente
    2) se REV manca, prova a ricavarlo dal suffisso finale del codice documento
       (es. 'RMS-C-ESE-COR-URB-001-00' -> 'Rev_00')
    """
    rev_s = safe_str(rev)
    code_s = safe_str(code)

    if rev_s:
        m = re.search(r"(\d{1,2})$", rev_s)
        if rev_s.isdigit() or m:
            num = int(m.group(1) if m else rev_s)
            return f"Rev_{num:02d}"
        if re.match(r"^rev[_\- ]?\d{1,2}$", rev_s, flags=re.IGNORECASE):
            m2 = re.search(r"(\d{1,2})$", rev_s)
            if m2:
                return f"Rev_{int(m2.group(1)):02d}"
        return rev_s

    m_code = re.search(r"([_-])(\d{2})$", code_s)
    if m_code:
        return f"Rev_{m_code.group(2)}"

    return ""

def sanitize_filename(s: str, fallback: str = "FILE") -> str:
    s = safe_str(s)
    if not s:
        return fallback
    s = s.replace("\\", "_").replace("/", "_").replace(":", "_")
    s = re.sub(r"[<>\"|?*]", "_", s)
    s = re.sub(r"\s+", "", s)
    return s or fallback


def _set_cell_text_preserve_format(cell, new_text: str) -> None:
    """Imposta il testo di una cella di tabella preservando, per quanto possibile,
    la formattazione del primo run (font, bold, size). Se la cella ha piu
    paragrafi, rimuove quelli successivi al primo.
    Usato per sovrascrivere la data di emissione nella tabella revisioni
    senza perdere lo stile grafico del template.
    """
    try:
        if not cell.paragraphs:
            cell.text = safe_str(new_text)
            return
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = safe_str(new_text)
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = safe_str(new_text)
        # rimuovi paragrafi extra nella cella
        for extra_p in list(cell.paragraphs[1:]):
            el = extra_p._element
            if el.getparent() is not None:
                el.getparent().remove(el)
    except Exception:
        try:
            cell.text = safe_str(new_text)
        except Exception:
            pass

def safe_read_excel(path: str) -> pd.DataFrame:
    df = pd.read_excel(path)
    df.columns = (
        df.columns.astype(str)
        .str.strip()
        .str.replace(r"[^\w]", "_", regex=True)
        .str.replace("__+", "_", regex=True)
    )
    return df


def _normalize_df_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Applica la stessa normalizzazione colonne di safe_read_excel."""
    df.columns = (
        df.columns.astype(str)
        .str.strip()
        .str.replace(r"[^\w]", "_", regex=True)
        .str.replace("__+", "_", regex=True)
    )
    return df


def read_report_smart(path: str) -> pd.DataFrame:
    """Legge dal file Report il foglio più utile per il lookup titoli.

    Strategia: scorre tutti i fogli del workbook e assegna un punteggio
    basato sulla presenza di colonne 'codice' e 'titolo'. Restituisce il
    foglio con punteggio più alto. In questo modo funziona anche con
    Report_Completo.xlsx dove il foglio utile è 'Verifica Elaborati'
    (non il primo foglio 'Nota Ricezione').
    """
    if not path:
        return pd.DataFrame()
    try:
        xl = pd.ExcelFile(path)
    except Exception:
        return pd.DataFrame()

    best_df = None
    best_score = -1

    for sheet in xl.sheet_names:
        try:
            df = pd.read_excel(path, sheet_name=sheet)
            df = _normalize_df_columns(df)
            score = 0
            for c in df.columns:
                lc = c.lower()
                if "codice" in lc and "unnamed" not in lc:
                    score += 2
                if "titolo" in lc and "unnamed" not in lc:
                    score += 2
                if "revis" in lc or "_rev" in lc:
                    score += 1
            # penalizza fogli con troppo poche righe utili
            non_empty = df.dropna(how="all")
            if len(non_empty) < 3:
                score -= 5
            if score > best_score:
                best_score = score
                best_df = df
        except Exception:
            continue

    return best_df if best_df is not None else pd.DataFrame()


# ---------------------------------------------------------------------------
# Lettura "Gruppi disciplinari - ispettori" (Excel)
# ---------------------------------------------------------------------------
def _find_header_row_for_gruppi(raw: pd.DataFrame) -> int | None:
    """Find header row index in a 'Gruppi disciplinari-ispettori' Excel export.

    We look for a row containing at least 'Titolo dell'elaborato' and 'Argomento'
    (case-insensitive, tolerant to newlines).
    """
    if raw is None or raw.empty:
        return None

    def norm_cell(x: object) -> str:
        return normalize_text(str(x)).replace(" ", "")

    target_a = "TITOLODELLELABORATO"
    target_b = "ARGOMENTO"

    for i in range(min(len(raw), 50)):  # scan first rows only
        row = raw.iloc[i].tolist()
        joined = " ".join(norm_cell(v) for v in row)
        if target_a in joined and target_b in joined:
            return i
    return None

def read_gruppi_disciplinari_excel(path: str, disciplina_col_name: str = "") -> pd.DataFrame:
    """Read the 'Gruppi disciplinari-ispettori' Excel (e.g. CORMOR-IV_...xlsx).

    Returns a normalized dataframe with at least:
    - Titolo_dell_elaborato
    - Codice_Tecne_LINK
    - Codice_ASPI
    - Doc
    - Argomento
    - ISP
    """
    if not path:
        return pd.DataFrame()

    # First pass: raw (no header) to detect header row
    raw = pd.read_excel(path, sheet_name=0, header=None)
    hdr = _find_header_row_for_gruppi(raw)
    if hdr is None:
        # fallback: try header=0
        df = pd.read_excel(path, sheet_name=0)
    else:
        df = pd.read_excel(path, sheet_name=0, header=hdr)

    # Normalize column names (keep meaning)
    df.columns = [str(c).strip().replace("\n", " ") for c in df.columns]

    # Drop fully empty rows
    df = df.dropna(how="all").copy()

    # Standardize expected columns (tolerant matching)
    def pick_col(candidates: list[str]) -> str | None:
        for c in df.columns:
            nc = normalize_text(c)
            for cand in candidates:
                if cand in nc:
                    return c
        return None

    col_tit = pick_col(["titolo dell'elaborato", "titolo elaborato", "titolo"])
    col_tec = pick_col(["codice tecne", "codice tecne (link)", "codice tecne link", "codice tecne"])
    col_aspi = pick_col(["codice aspi"])
    col_doc = pick_col(["doc."])
    col_arg = pick_col(["argomento"])

    # Colonna disciplina/gruppo: può variare. L'utente può indicarla in maschera (es. "Argomento").
    col_disc = None
    if disciplina_col_name:
        # match case-insensitive contro i nomi colonna reali
        for c in df.columns:
            if normalize_text(c) == normalize_text(disciplina_col_name):
                col_disc = c
                break
    if col_disc is None:
        # fallback automatico
        col_disc = pick_col(["argomento", "disciplina", "gruppo", "gruppi disciplinari", "gruppo disciplinare"])
    # se non trova nulla, usa comunque Argomento (se presente)
    if col_disc is None:
        col_disc = col_arg

    col_isp = pick_col(["isp"])

    out = pd.DataFrame()
    if col_tit: out["Titolo_dell_elaborato"] = df[col_tit]
    if col_tec: out["Codice_Tecne_LINK"] = df[col_tec]
    if col_aspi: out["Codice_ASPI"] = df[col_aspi]
    if col_doc: out["Doc"] = df[col_doc]
    if col_arg: out["Argomento"] = df[col_arg]
    if col_disc: out["DISCIPLINA_SRC"] = df[col_disc]
    if col_isp: out["ISP"] = df[col_isp]

    # Clean strings
    for c in out.columns:
        out[c] = out[c].apply(lambda x: safe_str(x).strip())

    # Remove group header rows (often with empty codes)
    # Keep rows that have at least a title + (one code among tec/aspi/doc) or an argomento
    def _row_ok(r):
        has_title = bool(safe_str(r.get("Titolo_dell_elaborato", "")).strip())
        has_code = bool(safe_str(r.get("Codice_Tecne_LINK", "")).strip() or safe_str(r.get("Codice_ASPI", "")).strip() or safe_str(r.get("Doc", "")).strip())
        has_arg = bool(safe_str(r.get("Argomento", "")).strip())
        return has_title and (has_code or has_arg)

    if not out.empty:
        out = out[out.apply(_row_ok, axis=1)].copy()

    return out

def _normalize_sp(s: str) -> str:
    s = safe_str(s).strip()
    s = s.replace("_", "-")
    # remove trailing "-0" / "_0"
    if s.endswith("-0"):
        s = s[:-2]
    return s.upper()

def build_doc_to_isp_map(gruppi_df: pd.DataFrame) -> dict:
    """Build a map 'normalized_doc_code' -> 'ISP' from gruppi_df."""
    mappa = {}
    if gruppi_df is None or gruppi_df.empty:
        return mappa

    for _, r in gruppi_df.iterrows():
        isp = safe_str(r.get("ISP", "")).strip()
        if not isp:
            continue

        for key in ["Doc", "Codice_ASPI", "Codice_Tecne_LINK"]:
            val = safe_str(r.get(key, "")).strip()
            if not val:
                continue
            mappa[normalize_code(val)] = isp
            # also store SP-normalized variant if it looks like SP code
            mappa[normalize_code(_normalize_sp(val))] = isp
    return mappa


def find_ispettore_for_doc(doc_code: str, doc_to_isp: dict) -> str:
    """Trova l'ispettore per un codice documento usando la mappa doc_to_isp.
    Applica normalizzazioni per gestire varianti (SP con/ senza -0, underscore, ecc.)."""
    if not doc_to_isp or not doc_code:
        return ""
    candidates = []
    dc = safe_str(doc_code).strip()
    if dc:
        candidates.append(dc)
        candidates.append(normalize_code(dc))
        candidates.append(_normalize_sp(dc))
        candidates.append(normalize_code(_normalize_sp(dc)))
    for c in candidates:
        k = normalize_code(c)
        if k in doc_to_isp:
            return safe_str(doc_to_isp.get(k, "")).strip()
    return ""

def docs_from_gruppi_for_disciplina(gruppi_df: pd.DataFrame, disciplina: str, codice_sp_disc: str) -> list[dict]:
    """Return list of docs (codice, rev, titolo) for the given disciplina using gruppi_df.

    Match priority:
    A) Argomento ~ disciplina (case/space insensitive containment)
    B) Doc ~ codice_sp_disc (SP code, tolerant to trailing '-0')
    """
    if gruppi_df is None or gruppi_df.empty:
        return []

    kd = normalize_text(disciplina)
    kd_comp = re.sub(r"\s+", "", kd)

    sp_norm = _normalize_sp(codice_sp_disc)

    def _match_arg(v: str) -> bool:
        kv = normalize_text(v)
        kv_comp = re.sub(r"\s+", "", kv)
        return bool(kd_comp) and (kd_comp in kv_comp or kv_comp in kd_comp)

    df = gruppi_df.copy()

    if "DISCIPLINA_SRC" in df.columns or "Argomento" in df.columns:
        col_disc_use = "DISCIPLINA_SRC" if "DISCIPLINA_SRC" in df.columns else "Argomento"
        df_a = df[df[col_disc_use].apply(_match_arg)]
    else:
        df_a = df.iloc[0:0]

    if df_a.empty and "Doc" in df.columns:
        df_b = df[df["Doc"].apply(lambda x: _normalize_sp(x) == sp_norm)]
    else:
        df_b = df_a

    df_use = df_b

    docs = []
    for _, r in df_use.iterrows():
        titolo = safe_str(r.get("Titolo_dell_elaborato", "")).strip()
        codice = (safe_str(r.get("Codice_Tecne_LINK", "")).strip() or
                  safe_str(r.get("Codice_ASPI", "")).strip() or
                  safe_str(r.get("Doc", "")).strip())
        if not titolo or not codice:
            continue
        docs.append({"codice": codice, "rev": safe_str(r.get("Rev", "")).strip(), "titolo": titolo})

    return unique_docs(docs)

def traduci_stato(tag: str, status: str) -> str:
    t = normalize_text(tag + " " + status)
    if "CLOSED" in t:
        return "Chiusa"
    if "WAITING" in t:
        return "In attesa"
    return "Aperta"

def is_nessun_rilievo(tag: str, *other_parts) -> bool:
    """Restituisce True SOLO se il campo TAG indica assenza di rilievi.
    Non controlla title/description per evitare falsi positivi (un NC può
    citare "nessun rilievo" nel testo della descrizione).
    Se il tag è vuoto o non riconosciuto, controlla anche title (primo elemento
    di other_parts) ma NON la description.
    """
    tag_norm = normalize_text(safe_str(tag))
    if "NESSUN RILIEVO" in tag_norm:
        return True
    # Tag non riconosciuto (vuoto o generico) → controlla anche il titolo
    # ma SOLO se il tag non è NC/OSS
    if not tag_norm or (tag_norm not in ("NC", "OSS") and "NC" not in tag_norm and "OSS" not in tag_norm):
        title_norm = normalize_text(safe_str(other_parts[0])) if other_parts else ""
        if "NESSUN RILIEVO" in title_norm:
            return True
    return False

def extract_document_code_from_title(title: str) -> str:
    """
    Estrae un codice documento dal titolo ToDo, se contiene stringhe tipo:
    PV032-PE-PEST-GEO-00000-REL-000001_00
    oppure T1462-...-00 ecc.
    """
    t = safe_str(title)
    if not t:
        return ""
    # cerca token con molti blocchi separati da - o _
    candidates = re.findall(r"[A-Z0-9]{2,}(?:[-_][A-Z0-9]{1,}){4,}", t.upper())
    if not candidates:
        return ""
    # prendi il più lungo
    cand = max(candidates, key=len)
    return cand.replace("_", "-")  # normalizziamo visivamente


def extract_all_document_codes_from_title(title: str) -> list:
    """Estrae TUTTI i codici documento presenti nel titolo ToDo.

    Gestisce i casi in cui un singolo ToDo copre piu elaborati scritti nello
    stesso campo Title, separati da ';' , ',' o newline, es.:
        'RMS_G_ESE_COR_ARC_030_00; RMS_G_ESE_COR_ARC_032_00'
    -> ['RMS-G-ESE-COR-ARC-030-00', 'RMS-G-ESE-COR-ARC-032-00']

    Ritorna lista vuota se il titolo non contiene alcun codice documento
    riconoscibile (es. 'RILIEVO GENERALE' o nomi cartella come
    'RE00_E- ARCHITETTURA STATO DI FATTO').
    """
    t = safe_str(title)
    if not t:
        return []
    # split su separatori tipici per multi-codice
    parts = re.split(r"[;,\n\r]+", t)
    out = []
    seen = set()
    for p in parts:
        p = p.strip()
        if not p:
            continue
        candidates = re.findall(r"[A-Z0-9]{2,}(?:[-_][A-Z0-9]{1,}){4,}", p.upper())
        if candidates:
            # prendi il codice piu lungo presente nel segmento
            cand = max(candidates, key=len).replace("_", "-")
            if cand not in seen:
                seen.add(cand)
                out.append(cand)
    return out


def docs_from_report_for_folder(report_df: pd.DataFrame, folder_title: str) -> list[dict]:
    """Ricava gli elaborati dal Report a partire da un nome cartella/gruppo.

    Esempio ToDo: "RE00_E- ARCHITETTURA STATO DI FATTO".
    Nel report i PDF reali sono elencati sotto la colonna "Cartella"
    (es. "E ARCHITETTURA STATO DI FATTO").
    Restituisce i singoli elaborati contenuti nella cartella, evitando che
    il nome cartella finisca come pseudo-codice nel riepilogo finale.
    """
    if report_df is None or report_df.empty:
        return []

    cols = {str(c).strip().lower(): c for c in report_df.columns}
    col_folder = next((c for c in report_df.columns if 'cartella' in str(c).lower()), None)
    col_code = next((c for c in report_df.columns if 'codice elaborato' in str(c).lower() or str(c).lower() == 'codice_elaborato'), None)
    col_title = next((c for c in report_df.columns if 'titolo elenco' in str(c).lower() or 'titolo' in str(c).lower()), None)
    if not col_folder or not col_code:
        return []

    raw = safe_str(folder_title)
    if not raw:
        return []
    # elimina eventuale prefisso tecnico tipo RE00_E- / RE00_A- ecc.
    folder_key = re.sub(r'^[A-Z]{2}\d{2}[_-][A-Z]-\s*', '', raw, flags=re.IGNORECASE)
    folder_key_norm = normalize_text(folder_key)
    folder_key_comp = re.sub(r'\s+', '', folder_key_norm)
    if not folder_key_comp:
        return []

    out = []
    for _, rr in report_df.iterrows():
        folder_val = safe_str(rr.get(col_folder, ''))
        if not folder_val:
            continue
        f_norm = normalize_text(folder_val)
        f_comp = re.sub(r'\s+', '', f_norm)
        if folder_key_comp not in f_comp and f_comp not in folder_key_comp:
            continue
        code = safe_str(rr.get(col_code, ''))
        if not code:
            continue
        out.append({
            'code': code,
            'rev': '',
            'titolo': safe_str(rr.get(col_title, '')) if col_title else '',
        })

    # de-dup preserving order
    seen = set()
    uniq = []
    for d in out:
        nk = normalize_code(d.get('code', ''))
        if not nk or nk in seen:
            continue
        seen.add(nk)
        uniq.append(d)
    return uniq


# --- INIZIALI ISPETTORI ------------------------------------------------------
def build_initials_map_from_elenco(elenco_df, col_red):
    """Crea una mappa {cognome_upper -> iniziali} leggendo il campo
    Nome redattore di ELENCO_ELABORATI.

    Formato atteso delle celle:
        'Arch. S. Arcangelelli (SA); P. I. M. Garofalo (MG); Ing. O. Bellaroba (OB)'
    Estrae tutte le coppie (nome_parte, iniziali) e registra il cognome
    (ultima parola capitalizzata del nome_parte) come chiave.
    """
    m = {}
    if elenco_df is None or col_red is None or col_red not in elenco_df.columns:
        return m
    for val in elenco_df[col_red].dropna():
        s = str(val)
        for match in re.finditer(r"([^;\n(]+?)\s*\(([A-Z]{2,4})\)", s):
            name_part = match.group(1).strip()
            initials = match.group(2).strip()
            # prendi cognome: ultima parola capitalizzata (almeno 3 lettere) del nome
            words = re.findall(r"\b([A-Z][a-z]+)", name_part)
            # filtra titoli (Arch, Ing, Geom, Dott, Prof) che non sono cognomi
            titles = {"ARCH", "ING", "GEOM", "DOTT", "PROF", "AVV"}
            words_no_titles = [w for w in words if w.upper() not in titles]
            if words_no_titles:
                surname = words_no_titles[-1].upper()
                m.setdefault(surname, initials)
            elif words:
                m.setdefault(words[-1].upper(), initials)
    return m


def to_inspector_initials(full_name: str, initials_map: dict,
                          alias_map: dict = None) -> str:
    """Converte un nome esteso (da ToDo 'Created by' o da ELENCO)
    nelle iniziali da stampare in colonna ISPETTORE.

    Logica (in ordine):
      1) Se il nome contiene '(XX)' in coda -> restituisce XX.
      2) Applica eventuali alias (_ISPETTORE_ALIAS) per sostituire il nome.
         Dopo l'alias ricontrolla '(XX)' in coda.
      3) Cerca per cognome (ultima parola capitalizzata) nella initials_map.
      4) Fallback: iniziali dal primo e ultimo nome capitalizzato
         (es. 'Stefano Arcangelelli' -> 'SA').
      5) Se nulla di tutto cio funziona, restituisce il nome originale.
    """
    s = safe_str(full_name).strip()
    if not s:
        return ""

    # 1) iniziali gia presenti come '(XX)' in coda
    m = re.search(r"\(([A-Z]{2,4})\)\s*$", s)
    if m:
        return m.group(1)

    # 2) applica alias (nome esteso -> altro nome esteso)
    s_alias = s
    if alias_map:
        for src, dst in alias_map.items():
            if normalize_text(src) in normalize_text(s):
                s_alias = dst
                m2 = re.search(r"\(([A-Z]{2,4})\)\s*$", s_alias)
                if m2:
                    return m2.group(1)
                break

    # 3) match per cognome nella initials_map (prova dall'ultima parola)
    words = re.findall(r"\b([A-Z][a-z]+)", s_alias)
    titles = {"ARCH", "ING", "GEOM", "DOTT", "PROF", "AVV"}
    for w in reversed(words):
        if w.upper() in titles:
            continue
        if len(w) >= 3 and w.upper() in initials_map:
            return initials_map[w.upper()]

    # 4) fallback: primo + ultimo capitalizzato (escludendo titoli)
    non_title_words = [w for w in words if w.upper() not in titles]
    if len(non_title_words) >= 2:
        return (non_title_words[0][0] + non_title_words[-1][0]).upper()
    if non_title_words:
        return non_title_words[0][:2].upper()
    if words:
        return words[0][:2].upper()

    # 5) ultima ratio
    return s


# --- PARSING DATA DA GUI -----------------------------------------------------
def normalize_gui_date(s: str) -> str:
    """Normalizza una data inserita in GUI.
    Accetta gg/mm/aaaa, gg.mm.aaaa, gg-mm-aaaa (anche con anno a 2 cifre).
    Ritorna la data nel formato 'gg.mm.aaaa' (coerente col template).
    Se il formato non e riconosciuto, ritorna la stringa originale strippata.
    """
    t = safe_str(s).strip()
    if not t:
        return ""
    for fmt in ("%d/%m/%Y", "%d.%m.%Y", "%d-%m-%Y",
                "%d/%m/%y", "%d.%m.%y", "%d-%m-%y"):
        try:
            d = datetime.datetime.strptime(t, fmt).date()
            return d.strftime("%d.%m.%Y")
        except Exception:
            pass
    return t


# ============================================================
# LOOKUP TITOLI (Report coerenze / non coerenze)
# ============================================================

def _normalize_report_colname(s: str) -> str:
    # Normalizza intestazioni Excel per matching robusto
    return re.sub(r"\s+", " ", safe_str(s)).strip().lower()

def _report_code_variants(raw_code: str) -> list[str]:
    """
    Genera varianti chiave per match titoli dal Report:
    - strict: codice completo normalizzato
    - flex: rimuove il blocco a 4 cifre dopo T####- (es. T1224-0000 vs T1224-0001)
    - suffix: ignora completamente il prefisso T####-0000- (match sulla parte PE-...-00)
    Tutte le varianti sono in formato normalize_code (underscore).
    """
    s = safe_str(raw_code).strip().upper().replace("_", "-")
    if not s:
        return []
    strict = normalize_code(s)

    # flex: T####_0000_...  -> T####_...
    flex = strict
    flex = re.sub(r"^(T\d{4})_\d{4}_", r"\1_", flex)

    # suffix: rimuove T####_0000_ e mantiene il resto
    suffix = re.sub(r"^T\d{4}_\d{4}_", "", strict)

    # Ulteriore: se il codice nel report contiene spazi strani o doppie sezioni, pulisci multipli '_'
    variants = []
    for v in (strict, flex, suffix):
        v = re.sub(r"_+", "_", safe_str(v)).strip("_")
        if v and v not in variants:
            variants.append(v)
    return variants

def build_report_lookup(df: pd.DataFrame) -> dict:
    """
    Crea lookup dal Report non coerenze per ricavare:
    - Titolo Documento
    - Revisione
    Chiavi robuste: strict + varianti (flex + suffix).
    """
    # priorità assoluta alle intestazioni note
    col_map = { _normalize_report_colname(c): c for c in df.columns }

    col_codice = None
    # Nome che mi hai dato: "Codice Documento (da Elenco elaborati)"
    for k, orig in col_map.items():
        if "codice documento (da elenco elaborati)" in k:
            col_codice = orig
            break
    if col_codice is None:
        # fallback: prima colonna che contiene 'codice documento'
        for k, orig in col_map.items():
            if "codice documento" in k and "elenco" in k:
                col_codice = orig
                break
    if col_codice is None:
        # fallback generico
        col_codice = next((c for c in df.columns if "codice" in c.lower()), None)

    col_titolo = None
    for k, orig in col_map.items():
        if k == "titolo documento" or "titolo documento" in k:
            col_titolo = orig
            break
    if col_titolo is None:
        col_titolo = next((c for c in df.columns if "titolo" in c.lower()), None)

    col_rev = next((c for c in df.columns if "revis" in c.lower()), None)

    lookup = {}
    for _, r in df.iterrows():
        raw_code = safe_str(r.get(col_codice, "")) if col_codice else ""
        if not raw_code:
            continue

        titolo = safe_str(r.get(col_titolo, "")) if col_titolo else ""
        rev = safe_str(r.get(col_rev, "")) if col_rev else ""

        if rev.isdigit():
            rev = f"Rev_{int(rev):02d}"
        elif rev and not rev.upper().startswith("REV"):
            rev = f"Rev_{rev}"

        # Inserisci tutte le varianti: se collisione, preferisci la voce con titolo non vuoto
        for key in _report_code_variants(raw_code):
            if key not in lookup:
                lookup[key] = {"titolo": titolo, "revisione": rev, "raw": raw_code}
            else:
                # aggiorna se prima era vuoto e ora abbiamo titolo
                if (not safe_str(lookup[key].get("titolo", ""))) and titolo:
                    lookup[key] = {"titolo": titolo, "revisione": rev, "raw": raw_code}
    return lookup


def build_elenco_lookup(df: pd.DataFrame) -> dict:
    """
    Costruisce un lookup (codice -> titolo) direttamente dall'ELENCO ELABORATI.
    È il fallback più affidabile quando il report coerenze/non-coerenze non contiene tutti i codici.

    NOTA: non usa fallback generici (es. 'CODICE_COMMESSA') per evitare di mappare codici
    commessa/progetto al posto di codici documento. Se la colonna non è trovata, restituisce
    lookup vuoto e il sistema userà correttamente il report_lookup.
    """
    # Cerca colonna codice: deve contenere 'elaborat' O 'documento' nel nome
    col_cod = None
    for c in df.columns:
        lc = c.lower()
        if "codice" in lc and ("elaborat" in lc or "documento" in lc):
            col_cod = c
            break
    # NESSUN fallback generico: evita di prendere CODICE_COMMESSA o simili
    # Se l'elenco non ha colonne di codici elaborato, il lookup resta vuoto.

    col_tit = None
    for c in df.columns:
        lc = c.lower()
        if "titolo" in lc and ("elaborat" in lc or "document" in lc):
            col_tit = c
            break
    if col_tit is None:
        col_tit = next((c for c in df.columns if "titolo" in c.lower()), None)

    lookup = {}
    if not col_cod:
        return lookup

    for _, r in df.iterrows():
        raw_code = safe_str(r.get(col_cod, ""))
        if not raw_code:
            continue
        key = normalize_code(raw_code)
        titolo = safe_str(r.get(col_tit, "")) if col_tit else ""
        lookup[key] = {"titolo": titolo, "raw": raw_code}
    return lookup


def find_titolo_elaborato(todo_title: str, report_lookup: dict, elenco_lookup: dict) -> tuple[str, str]:
    """
    Restituisce (Titolo Elaborato, Revisione) cercando:
    1) nel report (titolo + revisione)
    2) fallback nell'elenco elaborati (solo titolo)
    """
    titolo, rev = find_in_lookup(report_lookup, todo_title)
    if titolo:
        return titolo, rev

    # fallback elenco
    title = safe_str(todo_title)
    norm = normalize_code(title)
    base = strip_revision_suffix(norm)

    if norm in elenco_lookup:
        return elenco_lookup[norm]["titolo"], ""
    if base in elenco_lookup:
        return elenco_lookup[base]["titolo"], ""

    code_tok = extract_document_code_from_title(title)
    if code_tok:
        norm2 = normalize_code(code_tok)
        base2 = strip_revision_suffix(norm2)
        if norm2 in elenco_lookup:
            return elenco_lookup[norm2]["titolo"], ""
        if base2 in elenco_lookup:
            return elenco_lookup[base2]["titolo"], ""

    # containment fallback
    for k in elenco_lookup:
        if base and (base in k or k in base):
            return elenco_lookup[k]["titolo"], ""

    return "", ""


def find_in_lookup(lookup: dict, todo_title: str) -> tuple[str, str]:
    """
    Restituisce (Titolo Elaborato, Revisione) a partire dal campo Title del ToDo.
    Strategia:
    1) prova match diretto normalizzato
    2) prova rimuovendo suffix _00
    3) prova estraendo un codice documento dal title
       - match strict + varianti (flex/suffix)
    4) fallback: match per contenimento (base in key o viceversa)
    """
    title = safe_str(todo_title)
    norm = normalize_code(title)
    base = strip_revision_suffix(norm)

    # helper: prova una lista di chiavi in ordine
    def _try(keys: list[str]) -> tuple[str, str]:
        for k in keys:
            if k in lookup:
                return safe_str(lookup[k].get("titolo", "")), safe_str(lookup[k].get("revisione", ""))
        return "", ""

    # 1-2) diretto
    tit, rev = _try([norm, base])
    if tit:
        return tit, rev

    # 3) estrai codice dal title e prova varianti report
    code_tok = extract_document_code_from_title(title)
    if code_tok:
        norm2 = normalize_code(code_tok)
        base2 = strip_revision_suffix(norm2)

        # varianti (coerenti con _report_code_variants)
        flex2 = re.sub(r"^(T\d{4})_\d{4}_", r"\1_", norm2)
        suffix2 = re.sub(r"^T\d{4}_\d{4}_", "", norm2)

        flex2b = strip_revision_suffix(flex2)
        suffix2b = strip_revision_suffix(suffix2)

        tit, rev = _try([norm2, base2, flex2, flex2b, suffix2, suffix2b])
        if tit:
            return tit, rev

    # 4) fallback per contenimento
    for k in lookup:
        if base and (base in k or k in base):
            return safe_str(lookup[k].get("titolo", "")), safe_str(lookup[k].get("revisione", ""))

    return "", ""


# ============================================================
# IMMAGINI: sorgenti + salvataggio/rename
# ============================================================
def discover_image_columns(todo_df: pd.DataFrame):
    keys = ("foto", "image", "img", "snapshot", "attachment", "allegat")
    cols = []
    for c in todo_df.columns:
        lc = c.lower()
        if any(k in lc for k in keys):
            cols.append(c)
    return cols

def split_many_images(cell_value: str):
    s = safe_str(cell_value)
    if not s:
        return []
    parts = re.split(r"[\n;,\|]+", s)
    return [p.strip() for p in parts if p.strip()]

def download_or_copy_image(src: str, dst_path: str):
    src = safe_str(src)
    if not src:
        return False
    # URL
    if re.match(r"^https?://", src, flags=re.IGNORECASE):
        if requests is None:
            return False
        try:
            r = requests.get(src, timeout=30)
            r.raise_for_status()
            with open(dst_path, "wb") as f:
                f.write(r.content)
            return True
        except Exception:
            return False
    # path locale
    if os.path.exists(src):
        try:
            shutil.copy2(src, dst_path)
            return True
        except Exception:
            return False
    return False

def save_todo_images(row: pd.Series, label: str, images_dir: str, image_cols: list):
    os.makedirs(images_dir, exist_ok=True)

    sources = []
    for c in image_cols:
        sources.extend(split_many_images(row.get(c, "")))

    # dedup preservando ordine
    seen = set()
    sources_unique = []
    for s in sources:
        if s not in seen:
            seen.add(s)
            sources_unique.append(s)

    if not sources_unique:
        return []

    label_fs = sanitize_filename(label, fallback="TODO")
    saved = []

    if len(sources_unique) == 1:
        dst = os.path.join(images_dir, f"{label_fs}_Foto.png")
        if download_or_copy_image(sources_unique[0], dst):
            saved.append(dst)
        return saved

    for i, s in enumerate(sources_unique, start=1):
        dst = os.path.join(images_dir, f"{label_fs}_Foto{i}.png")
        if download_or_copy_image(s, dst):
            saved.append(dst)
    return saved

def _sort_foto_name(name: str):
    base = os.path.basename(name)
    m = re.search(r"[_-]Foto(\d+)?\b", base, flags=re.IGNORECASE)
    if m:
        return int(m.group(1) or 0)
    return 9999

def find_images_in_tree(photo_root: str, label: str, disciplina: str | None = None):
    if not photo_root or not os.path.isdir(photo_root):
        return []
    exts = (".png", ".jpg", ".jpeg", ".webp", ".bmp")
    label_norm = label.strip()
    prefix1 = f"{label_norm.lower()}_foto"
    prefix2 = f"{label_norm.lower()}-foto"

    def scan_dir(d: str):
        out = []
        try:
            for fn in os.listdir(d):
                low = fn.lower()
                if not low.endswith(exts):
                    continue
                if low.startswith(prefix1) or low.startswith(prefix2):
                    out.append(os.path.join(d, fn))
        except Exception:
            pass
        return out

    matches = []
    if disciplina:
        disc_dir = os.path.join(photo_root, str(disciplina).strip())
        if os.path.isdir(disc_dir):
            matches.extend(scan_dir(disc_dir))

    matches.extend(scan_dir(photo_root))

    if not matches:
        for root, _, files in os.walk(photo_root):
            for fn in files:
                low = fn.lower()
                if not low.endswith(exts):
                    continue
                if low.startswith(prefix1) or low.startswith(prefix2):
                    matches.append(os.path.join(root, fn))

    matches = list(dict.fromkeys(matches))
    matches.sort(key=_sort_foto_name)
    return matches

def copy_and_rename_images(local_paths: list, images_dir: str, label: str):
    os.makedirs(images_dir, exist_ok=True)
    out = []
    if not local_paths:
        return out

    multiple = len(local_paths) > 1
    for i, src_path in enumerate(local_paths, start=1):
        ext = os.path.splitext(src_path)[1].lower() or ".png"
        if multiple:
            dst_name = f"{label}_Foto{i}{ext}"
        else:
            dst_name = f"{label}_Foto{ext}"
        dst_path = os.path.join(images_dir, dst_name)
        try:
            if os.path.abspath(src_path) != os.path.abspath(dst_path):
                shutil.copy2(src_path, dst_path)
            out.append(dst_path)
        except Exception:
            continue
    return out

def gather_todo_images(row: pd.Series, label: str, images_dir: str, image_cols: list, photo_dir: str = "", disciplina: str = ""):
    final_paths = []
    if image_cols:
        try:
            final_paths.extend(save_todo_images(row, label=label, images_dir=images_dir, image_cols=image_cols))
        except Exception:
            pass

    local_matches = find_images_in_tree(photo_dir, label, disciplina=disciplina)
    if local_matches:
        if final_paths:
            os.makedirs(images_dir, exist_ok=True)
            start_idx = len(final_paths) + 1
            for j, src_path in enumerate(local_matches, start=start_idx):
                ext = os.path.splitext(src_path)[1].lower() or ".png"
                dst_name = f"{label}_Foto{j}{ext}"
                dst_path = os.path.join(images_dir, dst_name)
                try:
                    shutil.copy2(src_path, dst_path)
                    final_paths.append(dst_path)
                except Exception:
                    continue
        else:
            final_paths.extend(copy_and_rename_images(local_matches, images_dir=images_dir, label=label))

    return final_paths


# ============================================================
# WORD: bookmark + link interni + appendice immagini
# ============================================================
def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int):
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), bookmark_name)

    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)

def add_internal_hyperlink(paragraph, anchor: str, text: str, font_pt: int | float = 10):
    """Crea un hyperlink interno (bookmark) compatibile con python-docx/lxml."""
    hyperlink = OxmlElement("w:hyperlink")
    # anchor interno (bookmark) con namespace corretto
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # font size (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(float(font_pt) * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(float(font_pt) * 2)))
    rPr.append(szCs)


    # stile hyperlink (se presente nel doc)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # underline + color (fallback se lo stile non è applicato dal template)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # font size (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(float(font_pt) * 2)))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(float(font_pt) * 2)))
    rPr.append(szCs)


    # stile hyperlink (se presente nel doc)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # underline + color (fallback se lo stile non è applicato dal template)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def clear_cell(cell):
    for p in list(cell.paragraphs):
        try:
            p._element.getparent().remove(p._element)
        except Exception:
            pass
    cell.add_paragraph("")

def write_code_cell(cell, main_text: str, sub_text: str = "", main_pt: int | float = 10, sub_pt: int | float = 8, hyperlink_anchor: str | None = None):
    """Scrive nella cella: prima riga (main_text) a main_pt, seconda riga (sub_text) a sub_pt.
    Se hyperlink_anchor è valorizzato, la prima riga diventa un link interno al bookmark.
    """
    clear_cell(cell)
    p = cell.paragraphs[0]
    if hyperlink_anchor:
        add_internal_hyperlink(p, hyperlink_anchor, main_text, font_pt=main_pt)
    else:
        r = p.add_run(main_text)
        r.font.size = Pt(main_pt)
    if sub_text:
        r2 = p.add_run("\n" + sub_text)
        r2.font.size = Pt(sub_pt)


def next_bookmark_id(doc: Document) -> int:
    mx = 1
    for el in doc.element.iter():
        if el.tag.endswith('bookmarkStart'):
            try:
                mx = max(mx, int(el.get(qn('w:id'))))
            except Exception:
                pass
    return mx + 1

def ensure_appendix(doc: Document) -> int:
    for p in doc.paragraphs:
        if (p.text or "").strip().upper() == "ALLEGATI IMMAGINI":
            return next_bookmark_id(doc)

    doc.add_page_break()
    h = doc.add_paragraph("ALLEGATI IMMAGINI")
    try:
        h.style = "Heading 1"
    except Exception:
        pass
    return next_bookmark_id(doc)

def add_images_appendix_and_links(doc: Document, tab_rilievi, code_to_images: dict, code_to_desc: dict, row_code_map: dict[int, str] | None = None):
    """
    - Trasforma la prima colonna delle righe con immagini in link interno verso la PRIMA immagine associata.
    - Aggiunge in fondo (appendice) le immagini:
        * UNA immagine per pagina (page break tra immagini)
        * immagine CENTRATA
        * UNA SOLA didascalia sotto l'immagine (centrata)
    """
    if not code_to_images:
        return

    # Pillow (opzionale) per scalare correttamente in base al rapporto d'aspetto
    try:
        from PIL import Image  # type: ignore
    except Exception:
        Image = None

    bm_id = ensure_appendix(doc)

    # 1) tabella: prima colonna -> link interno se ci sono immagini.
    #    Il testo visualizzato resta nel formato richiesto: "NC1" / "OSS2" + riga sotto "(TR-xxx)".
    for r_i in range(1, len(tab_rilievi.rows)):
        row = tab_rilievi.rows[r_i].cells
        if not row:
            continue

        # Recupera il codice interno (NC-<LABEL> / OSS-<LABEL>) dalla mappa costruita durante la generazione,
        # così non dipendiamo dal testo in cella.
        code = None
        if row_code_map and r_i in row_code_map:
            code = row_code_map[r_i]
        else:
            # Fallback: prova a ricostruire da testo tipo "NC1\n(TR-011)" o "OSS2\n(TR-099)"
            raw0 = (row[0].text or "").strip()
            m_label = re.search(r"\(([^)]+)\)", raw0)
            m_tipo = re.match(r"\s*(NC|OSS)", raw0, flags=re.IGNORECASE)
            if m_label and m_tipo:
                code = f"{m_tipo.group(1).upper()}-{m_label.group(1).strip()}"

        if not code:
            continue

        imgs = code_to_images.get(code, [])
        if not imgs:
            continue

        # Mantieni il testo richiesto, ma trasformalo in hyperlink interno verso la prima immagine.
        display_txt = (row[0].text or "").strip()
        lines = [l.strip() for l in display_txt.splitlines() if l.strip()]
        main_txt = lines[0] if lines else code
        sub_txt = lines[1] if len(lines) > 1 else ""

        bm_prefix = f"IMG_{code}".replace("-", "_")

        write_code_cell(row[0], main_txt, sub_txt, main_pt=10, sub_pt=8, hyperlink_anchor=f"{bm_prefix}_1")

    # 2) appendice: una immagine per pagina, centrata, con caption sotto
    doc.add_paragraph("")  # spazio dopo heading "ALLEGATI IMMAGINI"

    first_image = True

    # Limiti "ragionevoli" per A4 portrait (con margini standard). Lasciamo spazio alla didascalia.
    MAX_W_IN = 6.0     # larghezza max immagine
    MAX_H_IN = 7.7     # altezza max immagine (per stare in pagina + caption)

    def _scaled_width_for_image(img_path: str) -> float:
        """Ritorna la width in inches da passare a add_picture per stare dentro MAX_W/ MAX_H."""
        if Image is None:
            return MAX_W_IN  # fallback: stessa width per tutte

        try:
            im = Image.open(img_path)
            w_px, h_px = im.size
            dpi = None
            try:
                dpi = im.info.get("dpi", None)
            except Exception:
                dpi = None
            if isinstance(dpi, tuple) and dpi and dpi[0]:
                xdpi = float(dpi[0])
            else:
                xdpi = 96.0  # fallback comune Windows
            if xdpi <= 0:
                xdpi = 96.0

            w_in = w_px / xdpi
            h_in = h_px / xdpi

            scale = min(
                MAX_W_IN / w_in if w_in > 0 else 1.0,
                MAX_H_IN / h_in if h_in > 0 else 1.0,
                1.0,
            )
            return max(1.0, w_in * scale)
        except Exception:
            return MAX_W_IN

    def _code_sort_key(c: str):
        tag = 0 if c.startswith("NC-") else 1
        return (tag, c)

    for code in sorted(code_to_images.keys(), key=_code_sort_key):
        imgs = code_to_images.get(code, [])
        if not imgs:
            continue
        desc = (code_to_desc.get(code) or "").strip()
        bm_prefix = f"IMG_{code}".replace("-", "_")

        for idx, img in enumerate(imgs, start=1):
            if not first_image:
                doc.add_page_break()
            first_image = False

            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_bookmark(p_img, f"{bm_prefix}_{idx}", bm_id)
            bm_id += 1

            run = p_img.add_run()
            try:
                w_in = _scaled_width_for_image(str(img))
                run.add_picture(str(img), width=Inches(w_in))
            except Exception:
                run.add_text(f"[Immagine non inserita: {os.path.basename(img)}]")

            cap = f"{code} - {desc}".strip(" -")
            p_cap = doc.add_paragraph(cap)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p_cap.runs:
                r.italic = True

    doc.add_paragraph("")


# ─── RENDER TEMPLATE ROBUSTO (zip-level) ─────────────────────────────────────
# Word spezza i tag Jinja2 {{ var }} su più run XML in header/footer.
# docxtpl passa il raw XML a Jinja2 → TemplateSyntaxError.
# Soluzione in due stadi:
#   1) Pre-elabora il .docx come ZIP sostituendo {{ var }} negli XML
#      di header/footer direttamente (scanner carattere per carattere).
#   2) Monkey-patch di build_headers_footers_xml: restituisce i blob
#      già pre-elaborati senza ulteriore processing Jinja2.
# ─────────────────────────────────────────────────────────────────────────────
import zipfile as _zipfile_mod
import types as _types_mod


def _subst_jinja_raw_xml(xml_str: str, context: dict) -> str:
    """Sostituisce {{ var }} nel raw XML anche se spezzato su più XML run.

    Scansione carattere-per-carattere: quando trova '{{' cerca il '}}' di
    chiusura, estrae il testo strippando i tag XML intermedi per ottenere
    il nome della variabile, e sostituisce l'intero span con il valore.
    """
    result = []
    i = 0
    n = len(xml_str)
    while i < n:
        if xml_str[i:i+2] == '{{' and xml_str[i:i+3] != '{{{':
            j = xml_str.find('}}', i + 2)
            if j == -1:
                result.append(xml_str[i]); i += 1; continue
            # Estrai nome variabile rimuovendo tag XML intermedi
            var_name = re.sub(r'<[^>]+>', '', xml_str[i+2:j]).strip()
            var_name = re.sub(r'\s+', '', var_name)  # strip spazi interni (Word spezza variabili)
            if re.match(r'^\w+$', var_name) and var_name in context:
                value = (str(context[var_name])
                         .replace('&', '&amp;')
                         .replace('<', '&lt;')
                         .replace('>', '&gt;'))
                result.append(value)
                i = j + 2
            else:
                result.append(xml_str[i]); i += 1
        else:
            result.append(xml_str[i]); i += 1
    return ''.join(result)


def _preprocess_template_zip(template_path: str, context: dict, output_path: str) -> None:
    """Copia il .docx sostituendo le variabili Jinja2 negli XML header/footer."""
    with _zipfile_mod.ZipFile(template_path, 'r') as zin:
        with _zipfile_mod.ZipFile(output_path, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if re.search(r'^word/(header|footer)\d*\.xml$', item.filename):
                    try:
                        fixed = _subst_jinja_raw_xml(data.decode('utf-8'), context)
                        data = fixed.encode('utf-8')
                    except Exception:
                        pass
                zout.writestr(item, data)


def render_docx_template(template_path: str, context: dict, output_path: str) -> None:
    """Render robusto del template Word.

    Gestisce definitivamente i tag Jinja2 {{ }} spezzati in header/footer
    combinando due tecniche:
    1) Pre-processing ZIP: sostituisce le variabili direttamente negli XML
       di header/footer prima che docxtpl li veda.
    2) Monkey-patch di build_headers_footers_xml: bypassa Jinja2 per
       header/footer (già pre-elaborati) e restituisce i blob as-is.
    Questo garantisce che Jinja2 non veda mai il raw XML degli header/footer.
    """
    pre_path = output_path + "__pre__.docx"
    try:
        _preprocess_template_zip(template_path, context, pre_path)
        tpl = DocxTemplate(pre_path)

        # Monkey-patch: restituisce i blob header/footer pre-elaborati
        # senza far passare nulla per Jinja2.
        # *args/**kwargs garantisce compatibilità con qualsiasi versione di docxtpl:
        #   vecchia: build_headers_footers_xml(self, context, jinja_env)
        #   nuova:   build_headers_footers_xml(self, context, uri, jinja_env)
        def _passthrough_headers(self, *args, **kwargs):
            uri_str = " ".join(str(a) for a in args if isinstance(a, str)).lower()
            if uri_str and "footer" in uri_str:
                targets = ["footer_parts"]
            elif uri_str and "header" in uri_str:
                targets = ["header_parts"]
            else:
                targets = ["header_parts", "footer_parts"]
            for attr in targets:
                try:
                    parts = getattr(self.docx.part, attr, {})
                    for relKey, part in parts.items():
                        yield relKey, part.blob.decode("utf-8")
                except Exception:
                    pass

        tpl.build_headers_footers_xml = _types_mod.MethodType(
            _passthrough_headers, tpl)

        tpl.render(context)
        tpl.save(output_path)
    finally:
        try:
            os.remove(pre_path)
        except Exception:
            pass


# ─── MAPPA ISPETTORI PER DISCIPLINA ──────────────────────────────────────────
_ISPETTORE_ALIAS: dict = {
    "Clara Soliman":      "P. I. M. Garofalo",
    "Carlo Renda":        "Ing. O. Bellaroba",
    "Gianluca Biaggioli": "Ing. O. Bellaroba",
}

_ISPETTORE_DISCIPLINA_OVERRIDE: dict = {
    "IMPIANTI": ["Ing. O. Bellaroba (OB)"],
}


def apply_ispettore_override(isp_raw: str, disciplina: str) -> str:
    """Applica la mappa ispettori: override completo per disciplina o
    sostituzione nome-per-nome tramite _ISPETTORE_ALIAS.
    """
    disc_key = re.sub(r'\s+', '', normalize_text(disciplina))
    for k, override_list in _ISPETTORE_DISCIPLINA_OVERRIDE.items():
        if k.replace(' ', '') in disc_key or disc_key in k.replace(' ', ''):
            return '; '.join(override_list)
    result = safe_str(isp_raw).strip()
    for alias_src, alias_dst in _ISPETTORE_ALIAS.items():
        if normalize_text(alias_src) in normalize_text(result):
            result = alias_dst
            break
    return result


# ============================================================
# GUI

# ============================================================
class SchedeGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Generatore Schede Ispettive - ITS v8.6.7")
        self.root.geometry("820x640")

        self.files = {k: tk.StringVar() for k in ["elenco", "todo", "report", "template"]}
        # v8.6.5: Campi data opzionali (override)
        self.data_ricezione_override = tk.StringVar()
        self.data_emissione_override = tk.StringVar()

        # Logo ITS + crediti
        self.logo_img = None
        logo_path = resource_path("logo_its.png")
        # Fallback: se l'utente tiene il logo accanto all'EXE o nella cartella di lavoro
        if not os.path.exists(logo_path):
            cwd_logo = os.path.join(os.getcwd(), "logo_its.png")
            if os.path.exists(cwd_logo):
                logo_path = cwd_logo
        if os.path.exists(logo_path):
            try:
                self.logo_img = tk.PhotoImage(file=logo_path)
                tk.Label(self.root, image=self.logo_img).pack(pady=(10, 2))
            except Exception:
                self.logo_img = None

        tk.Label(self.root, text="Applicazione realizzata da Arch. Giuseppe Pizzi", fg="#333", font=("Segoe UI", 9)).pack(pady=(0, 10))

        frame = tk.Frame(self.root)
        frame.pack(pady=10)

        for lbl, key in [
            ("📘 ELENCO ELABORATI", "elenco"),
            ("📄 ToDo Trimble (export Excel)", "todo"),
            ("📗 Report coerenze elaborati", "report"),
            ("📂 Template Word", "template"),
        ]:
            row = tk.Frame(frame)
            row.pack(fill="x", pady=4)
            tk.Label(row, text=lbl, width=30, anchor="w").pack(side="left")
            tk.Entry(row, textvariable=self.files[key], width=58, state="readonly").pack(side="left")
            tk.Button(row, text="Sfoglia", command=lambda k=key: self.browse(k)).pack(side="left")

        # v8.6.5: Campi data opzionali (compilabili dal PM) -----------------
        date_frame = tk.LabelFrame(
            self.root,
            text=" Date (opzionali - se vuote, vengono usate quelle del file Excel / del template) ",
            font=("Segoe UI", 9, "bold"),
            padx=8, pady=6,
        )
        date_frame.pack(fill="x", padx=20, pady=(4, 6))

        row_d1 = tk.Frame(date_frame)
        row_d1.pack(fill="x", pady=2)
        tk.Label(row_d1, text="📅 Data ricezione elaborati (prima pagina)",
                 width=42, anchor="w").pack(side="left")
        tk.Entry(row_d1, textvariable=self.data_ricezione_override, width=18).pack(side="left")
        tk.Label(row_d1, text="  es. 28/11/2025", fg="#666", font=("Segoe UI", 8)).pack(side="left")

        row_d2 = tk.Frame(date_frame)
        row_d2.pack(fill="x", pady=2)
        tk.Label(row_d2, text="📅 Data emissione scheda ispettiva (tabella revisioni)",
                 width=42, anchor="w").pack(side="left")
        tk.Entry(row_d2, textvariable=self.data_emissione_override, width=18).pack(side="left")
        tk.Label(row_d2, text="  es. 13/01/2026", fg="#666", font=("Segoe UI", 8)).pack(side="left")

        tk.Label(
            self.root,
            text="Nota: se nel ToDo sono presenti colonne Immagini/Snapshot/Foto (path o URL), verranno salvate e inserite in appendice.",
            fg="#444",
        ).pack(pady=4)

        tk.Button(
            self.root,
            text="🚀 Genera Schede Ispettive",
            command=self.emissione_schede,
            bg="#107C10",
            fg="white",
            font=("Segoe UI", 11, "bold"),
            width=48,
            height=2,
        ).pack(pady=12)

    def browse(self, key):
        path = filedialog.askopenfilename(
                filetypes=[("Excel", "*.xlsx")] if key != "template" else [("Word", "*.docx")]
            )
        if path:
            self.files[key].set(path)

    def emissione_schede(self):
        try:
            elenco_df = safe_read_excel(self.files["elenco"].get())
            todo_df = safe_read_excel(self.files["todo"].get())
            report_df = read_report_smart(self.files["report"].get())
            gruppi_df = None  # v8.6.4: file gruppi rimosso
            template_path = self.files["template"].get()
        except Exception as e:
            messagebox.showerror("Errore", f"Impossibile leggere i file:\n{e}")
            return

        # v8.6.5: Date opzionali dalla GUI ---------------------------------
        # Se compilate, sovrascrivono:
        #  - Data_ricezione letta da ELENCO_ELABORATI (prima pagina di ogni scheda)
        #  - il testo della data di emissione nella tabella revisioni (Tabella 0)
        data_ricezione_override = normalize_gui_date(self.data_ricezione_override.get())
        data_emissione_override = normalize_gui_date(self.data_emissione_override.get())

        # progetto (da elenco)
        try:
            col_cod_sp = next(c for c in elenco_df.columns if "codice" in c.lower() and "sp" in c.lower())
            col_tit_proj = next(c for c in elenco_df.columns if "titolo" in c.lower() and "progetto" in c.lower())
        except StopIteration:
            messagebox.showerror("Errore", "Nel file ELENCO ELABORATI non trovo 'Codice SP' o 'Titolo progetto'.")
            return

        # Mappa Codice_SP per disciplina (ELENCO ELABORATI = database)
        # Esempio: 'Documenti generali' -> IT22079AR-034-SP-0001-0; 'Ambiente e vincoli' -> IT22079AR-034-SP-0002-0
        col_disc_elenco = next((c for c in elenco_df.columns if str(c).strip().lower() == "disciplina"), None)
        if not col_disc_elenco:
            col_disc_elenco = next((c for c in elenco_df.columns if "disciplina" in str(c).lower()), None)

        titolo_progetto = safe_str(elenco_df[col_tit_proj].iloc[0])

        disc_to_sp = {}
        if col_disc_elenco:
            for _, rr in elenco_df.iterrows():
                dname = safe_str(rr.get(col_disc_elenco, ""))
                sp_val = safe_str(rr.get(col_cod_sp, ""))
                if dname and sp_val:
                    disc_to_sp[normalize_text(dname)] = sp_val

        # fallback: se non troviamo nulla, usa il primo
        default_sp = safe_str(elenco_df[col_cod_sp].iloc[0])

        # === Redattori per disciplina (ELENCO ELABORATI: header "Nome redattore") ===
        col_red_elenco = next((c for c in elenco_df.columns if str(c).strip().lower() == "nome redattore"), None)
        if col_red_elenco is None:
            col_red_elenco = next((c for c in elenco_df.columns if "nome redattore" in str(c).lower() or "nome_redattore" in str(c).lower()), None)
        if col_red_elenco is None:
            # fallback: colonna K (indice 10)
            try:
                col_red_elenco = elenco_df.columns[10]
            except Exception:
                col_red_elenco = None

        disc_to_red = {}
        if col_disc_elenco and col_red_elenco:
            for _, rr in elenco_df.iterrows():
                dname = safe_str(rr.get(col_disc_elenco, ""))
                red_val = safe_str(rr.get(col_red_elenco, ""))
                if not dname or not red_val:
                    continue

                # split multipli in cella (separati da ';' o a capo) + unique preserving order
                parts = [p.strip() for p in re.split(r"[;\n\r]+", red_val) if p.strip()]
                if not parts:
                    continue

                key = normalize_text(dname)
                disc_to_red.setdefault(key, [])
                for p in parts:
                    if p not in disc_to_red[key]:
                        disc_to_red[key].append(p)

        # fallback vuoto: NON vogliamo mai inserire "tutti" se manca match

        # === Lookup redattori per Codice_SP (più robusto del match per disciplina) ===
        sp_to_red = {}
        if col_cod_sp and col_red_elenco:
            for _, rr in elenco_df.iterrows():
                sp_val = safe_str(rr.get(col_cod_sp, ""))
                red_val = safe_str(rr.get(col_red_elenco, ""))
                if not sp_val or not red_val:
                    continue
                parts = [p.strip() for p in re.split(r"[;\n\r]+", red_val) if p.strip()]
                if not parts:
                    continue
                sp_to_red.setdefault(sp_val, [])
                for p in parts:
                    if p not in sp_to_red[sp_val]:
                        sp_to_red[sp_val].append(p)



        col_nota = next((c for c in elenco_df.columns if "nota" in c.lower() and "ricezione" in c.lower()), None)
        col_data = next((c for c in elenco_df.columns if "data" in c.lower() and "ricezione" in c.lower()), None)
        col_red = next((c for c in elenco_df.columns if "redattore" in c.lower()), None)
        # Fase di progetto (es. "PROGETTO ESECUTIVO") – usata come variabile nel template
        col_fase = next((c for c in elenco_df.columns if "fase" in c.lower()), None)
        fase_progetto = safe_str(elenco_df[col_fase].iloc[0]) if col_fase else ""

        # Nota/Data ricezione: la NOTA può essere comune, ma la DATA deve essere presa per DISCIPLINA (colonna L: Data_ricezione)
        nota_ricezione = safe_str(elenco_df[col_nota].iloc[0]) if col_nota else ""

        # Default (fallback): prima riga
        data_ricezione_default = safe_date_str(elenco_df[col_data].iloc[0]) if col_data else ""

        # Mappe per disciplina / Codice_SP -> Data_ricezione (senza orario)
        disc_to_data = {}
        sp_to_data = {}
        if col_data:
            for _, rr in elenco_df.iterrows():
                dname = safe_str(rr.get(col_disc_elenco, "")) if col_disc_elenco else ""
                sp_val = safe_str(rr.get(col_cod_sp, ""))
                dt_val = rr.get(col_data, None)

                dt_str = safe_date_str(dt_val) if (dt_val is not None and not (hasattr(pd, "isna") and pd.isna(dt_val))) else ""
                if dt_str:
                    if dname:
                        kdisc = normalize_text(dname)
                        disc_to_data.setdefault(kdisc, dt_str)
                    if sp_val:
                        sp_to_data.setdefault(sp_val, dt_str)

        nome_redattore = safe_str(elenco_df[col_red].iloc[0]) if col_red else ""


        lookup = build_report_lookup(report_df)
        elenco_lookup = build_elenco_lookup(elenco_df)

        # v8.6.5: mappa iniziali ispettori (cognome_upper -> 'SA', 'MG', 'OB', ...)
        # estratta dalle parentesi di 'Nome redattore' in ELENCO_ELABORATI
        initials_map = build_initials_map_from_elenco(elenco_df, col_red_elenco)

        # colonne ToDo minime
        try:
            col_ass = next(c for c in todo_df.columns if "assignee" in c.lower())
            col_tit = next(c for c in todo_df.columns if "title" in c.lower())
            col_des = next(c for c in todo_df.columns if "desc" in c.lower())
            col_tag = next(c for c in todo_df.columns if "tag" in c.lower())
            col_sta = next(c for c in todo_df.columns if "status" in c.lower())
            col_lastmod = next((c for c in todo_df.columns if 'last modified by' in c.lower()), None)
        except StopIteration:
            messagebox.showerror("Errore", "Nel file ToDo non trovo colonne minime richieste (Assignee, Title, Description, Tag, Status).")
            return

        col_label = next((c for c in todo_df.columns if "label" in c.lower()), None)
        image_cols = discover_image_columns(todo_df)

        # Espansione per disciplina (Assignee multi)
        # Nota: in Assignee(s) Trimble può usare separatori ';' (es. "Documenti generali; Ambiente e vincoli").
        # In quel caso la stessa NC/OSS deve finire in entrambe le schede, non generare una nuova scheda "A; B".
        expanded = []
        for _, r in todo_df.iterrows():
            ass_raw = safe_str(r.get(col_ass, ""))
            disciplines = split_assignees(ass_raw)

            if not disciplines:
                rr = r.copy()
                rr["DISCIPLINA"] = "NON ASSEGNATO"
                expanded.append(rr)
                continue

            for d in disciplines:
                rr = r.copy()
                rr["DISCIPLINA"] = d.strip()
                expanded.append(rr)

        todo = pd.DataFrame(expanded)
        if todo.empty:
            messagebox.showwarning("Attenzione", "Nessun ToDo trovato nel file.")
            return

        out_dir = os.path.join(os.path.dirname(self.files["elenco"].get()), "Schede generate da ToDo")
        os.makedirs(out_dir, exist_ok=True)

        images_root = os.path.join(out_dir, "Immagini")
        os.makedirs(images_root, exist_ok=True)

        photo_dir = ""  # cartella foto rimossa dalla maschera (si usano solo immagini da ToDo)

        # Mappa globale: codice documento -> ispettore (da file "Gruppi disciplinari-ispettori")
        doc_to_isp = build_doc_to_isp_map(gruppi_df) if gruppi_df is not None else {}

        for disciplina in sorted(todo["DISCIPLINA"].unique()):
            subset = todo[todo["DISCIPLINA"] == disciplina]
            if subset.empty:
                continue

            # Codice SP specifico per disciplina (da ELENCO ELABORATI)


            codice_sp_disc = disc_to_sp.get(normalize_text(disciplina), default_sp)



            # Fallback robusto: se finisce su default_sp (spesso 0001) prova match per inclusione/alias


            if codice_sp_disc == default_sp:


                kdisc = normalize_text(disciplina)


                alias_map = {


                    "ECONOMICO": "DOCUMENTI ECONOMICI",


                    "DOCUMENTI ECONOMICI": "DOCUMENTI ECONOMICI",


                    "ECONOMICI": "DOCUMENTI ECONOMICI",


                }


                kprobe = alias_map.get(kdisc, kdisc)


                kprobe_comp = re.sub(r"\s+", "", kprobe)


                for kk, spv in disc_to_sp.items():


                    kk_comp = re.sub(r"\s+", "", kk)


                    if kprobe_comp and (kprobe_comp in kk_comp or kk_comp in kprobe_comp):


                        codice_sp_disc = spv


                        break



            # Nome redattore per scheda: priorità a Codice_SP (robusto), fallback per disciplina


            nome_redattore_disc = ""


            if codice_sp_disc in sp_to_red:


                nome_redattore_disc = ";\n".join(sp_to_red[codice_sp_disc])


            else:


                kdisc = normalize_text(disciplina)


                if kdisc in disc_to_red:


                    nome_redattore_disc = ";\n".join(disc_to_red[kdisc])


                else:


                    kdisc_comp = re.sub(r"\s+", "", kdisc)


                    for kk, vv in disc_to_red.items():


                        if kdisc_comp and (kdisc_comp in re.sub(r"\s+", "", kk) or re.sub(r"\s+", "", kk) in kdisc_comp):


                            nome_redattore_disc = ";\n".join(vv)


                            break




            # Data ricezione elaborati per DISCIPLINA (da ELENCO ELABORATI, colonna L: Data_ricezione)
            data_ricezione_disc = ""
            if codice_sp_disc in sp_to_data:
                data_ricezione_disc = sp_to_data.get(codice_sp_disc, "")
            if not data_ricezione_disc:
                kdisc = normalize_text(disciplina)
                if kdisc in disc_to_data:
                    data_ricezione_disc = disc_to_data.get(kdisc, "")
                else:
                    # fallback per inclusione/alias (es. ECONOMICO <-> DOCUMENTI ECONOMICI)
                    kdisc_comp = re.sub(r"\s+", "", kdisc)
                    for kk, vv in disc_to_data.items():
                        kk_comp = re.sub(r"\s+", "", kk)
                        if kdisc_comp and (kdisc_comp in kk_comp or kk_comp in kdisc_comp):
                            data_ricezione_disc = vv
                            break
            if not data_ricezione_disc:
                data_ricezione_disc = data_ricezione_default

            # v8.6.5: override data ricezione da GUI (se l'utente l'ha compilata)
            if data_ricezione_override:
                data_ricezione_disc = data_ricezione_override

            # Render template (zip-level fix per header/footer Jinja2 spezzati)
            tmp = os.path.join(out_dir, "_tmp.docx")
            render_docx_template(template_path, {
                "Codice_SP": codice_sp_disc,
                "Titolo_progetto": titolo_progetto,
                "DISCIPLINA": disciplina,
                "Nota_ricezione_elaborati": nota_ricezione,
                "Data_ricezione": data_ricezione_disc,
                "Nome_redattore": nome_redattore_disc,
                "Fase_di_progetto": fase_progetto,
                "Fase_progetto": fase_progetto,  # alias senza "di_"
            }, tmp)
            doc = Document(tmp)

            # v8.6.5: override data emissione nella tabella revisioni (Tabella 0)
            # Il template contiene una riga tipo: '0 | 13.01.2026 | Prima Emissione | Ing. G. Biaggioli'
            # Se l'utente ha compilato il campo in GUI, sostituiamo la colonna 'Data'.
            if data_emissione_override:
                try:
                    rev_table = doc.tables[0]
                    date_like_re = re.compile(r"\d{1,2}[./-]\d{1,2}[./-]\d{2,4}")
                    target_row = None
                    for rr in rev_table.rows:
                        cells = rr.cells
                        if len(cells) >= 2:
                            txt0 = cells[0].text.strip()
                            txt1 = cells[1].text.strip()
                            # prima cella: codice revisione breve ('0', '1', 'A'...)
                            # seconda cella: deve contenere qualcosa che assomigli a una data
                            if re.match(r"^[0-9A-Za-z]{1,3}$", txt0) and date_like_re.search(txt1):
                                target_row = rr
                                break
                    if target_row is not None:
                        _set_cell_text_preserve_format(target_row.cells[1], data_emissione_override)
                except Exception:
                    pass

            try:
                tab_rilievi = doc.tables[-2]
                tab_riep = doc.tables[-1]
            except Exception:
                messagebox.showerror("Errore", "Nel template Word non trovo le tabelle attese (rilievi + riepilogo) in fondo al documento.")
                try:
                    os.remove(tmp)
                except Exception:
                    pass
                return

            nc_count = 0
            oss_count = 0

            # RIEPILOGO DOCUMENTI:
            # Deve includere TUTTI i documenti della disciplina (non solo quelli con NC/OSS).
            # Proviamo a ricavare l'elenco completo dal "Report coerenze elaborati" (preferibile).
            all_docs = []
            if gruppi_df is not None:
                all_docs = docs_from_gruppi_for_disciplina(gruppi_df, disciplina=disciplina, codice_sp_disc=codice_sp_disc)
            if not all_docs:
                all_docs = docs_from_report_for_disciplina(report_df, disciplina=disciplina, codice_sp_disc=codice_sp_disc)

            # Key: normalized code -> info
            riepilogo = {}
            for ddoc in all_docs:
                d_code = safe_str(ddoc.get("code", "") or ddoc.get("codice", ""))
                if not d_code:
                    continue
                d_norm = normalize_code(d_code)
                d_rev = safe_str(ddoc.get("rev", ""))
                d_tit = safe_str(ddoc.get("titolo", ""))

                # se manca il titolo nel report, prova a ricavarlo dal lookup (senza dipendere dal ToDo)
                if not d_tit:
                    d_tit2, d_rev2 = find_titolo_elaborato(d_code, report_lookup=lookup, elenco_lookup=elenco_lookup)
                    d_tit = d_tit or d_tit2
                    d_rev = d_rev or d_rev2

                riepilogo[d_norm] = {"code": d_code, "titolo": d_tit, "rev": d_rev, "NC": False, "OSS": False}

            # Se il report non consente di determinare i documenti per disciplina, il riepilogo verrà popolato
            # dai documenti incontrati nel ToDo (fallback), ma manterrà comunque la logica "ASSENZA NC/OSS".

            # ── PRE-PASSO: costruisci il riepilogo da TUTTE le righe ToDo ──────────
            # Flag separati:
            #   NC / OSS  = il documento ha almeno una NC o OSS
            #   NESSUN    = il documento ha almeno una riga "Nessun rilievo"
            # La colonna ASSENZA NC/OSS mostra X se NESSUN=True
            # (anche se il documento ha anche NC/OSS, l'assenza è visibile).
            #
            # v8.6.5:
            # - Titoli con piu codici separati da ';'/',' vengono SPEZZATI e ogni
            #   codice riceve il flag (es. 'RMS_G_ESE_COR_ARC_030_00; RMS_G_ESE_COR_ARC_032_00'
            #   flagga entrambi i documenti).
            # - Titoli senza alcun codice riconoscibile (nomi cartella come
            #   'RE00_E- ARCHITETTURA STATO DI FATTO') con tag "Nessun rilievo"
            #   vengono trattati come COMMENTO IN BLOCCO: il flag NESSUN si
            #   propaga a TUTTI i documenti della disciplina non ancora flaggati
            #   (ne NC ne OSS ne NESSUN). La voce cartella NON viene inserita
            #   nel riepilogo, cosi il conteggio resta coerente.
            # - Titoli generici (es. 'RILIEVO GENERALE') con tag NC/OSS: come
            #   prima, vengono aggiunti come voce "generica" (real_code=False)
            #   cosi i rilievi generali restano visibili nel riepilogo.
            for __ri, __rr in subset.iterrows():
                __title_p = safe_str(__rr.get(col_tit, ""))
                __tag_p   = safe_str(__rr.get(col_tag, ""))
                __is_ness = is_nessun_rilievo(__tag_p, __title_p)
                __tag_n_p = normalize_text(__tag_p)

                # Estrai tutti i codici documento presenti nel titolo
                __codes_list = extract_all_document_codes_from_title(__title_p)

                if __codes_list:
                    # CASO A: titolo con 1+ codici -> applica flag a ciascun codice
                    for __code_p in __codes_list:
                        __dn_p = normalize_code(__code_p)
                        __tit_p, __rev_p = find_titolo_elaborato(
                            __code_p, report_lookup=lookup, elenco_lookup=elenco_lookup)
                        if __dn_p not in riepilogo:
                            riepilogo[__dn_p] = {
                                "code": __code_p, "titolo": __tit_p, "rev": __rev_p,
                                "NC": False, "OSS": False, "NESSUN": False,
                                "real_code": True}
                        else:
                            if not riepilogo[__dn_p].get("titolo") and __tit_p:
                                riepilogo[__dn_p]["titolo"] = __tit_p
                            if not riepilogo[__dn_p].get("rev") and __rev_p:
                                riepilogo[__dn_p]["rev"] = __rev_p
                            # promuovi a real_code se prima era generico
                            riepilogo[__dn_p]["real_code"] = True
                        if __is_ness:
                            riepilogo[__dn_p]["NESSUN"] = True
                        elif "OSS" in __tag_n_p:
                            riepilogo[__dn_p]["OSS"] = True
                        else:
                            riepilogo[__dn_p]["NC"] = True

                elif __is_ness:
                    # CASO B: titolo senza codici riconoscibili + "Nessun rilievo"
                    # -> commento in blocco su cartella / gruppo di elaborati.
                    # 1) prova a espandere il nome cartella nei singoli elaborati
                    #    leggendo la colonna "Cartella" del Report.
                    # 2) in ogni caso propaga NESSUN a tutti i documenti reali della
                    #    disciplina non ancora flaggati.
                    __folder_docs = docs_from_report_for_folder(report_df, __title_p)
                    for __fd in __folder_docs:
                        __fcode = safe_str(__fd.get("code", ""))
                        if not __fcode:
                            continue
                        __fdn = normalize_code(__fcode)
                        __ftit = safe_str(__fd.get("titolo", ""))
                        if not __ftit:
                            __ftit2, __frev2 = find_titolo_elaborato(
                                __fcode, report_lookup=lookup, elenco_lookup=elenco_lookup)
                            __ftit = __ftit or __ftit2
                        if __fdn not in riepilogo:
                            riepilogo[__fdn] = {
                                "code": __fcode, "titolo": __ftit, "rev": "",
                                "NC": False, "OSS": False, "NESSUN": True,
                                "real_code": True}
                        else:
                            if not riepilogo[__fdn].get("titolo") and __ftit:
                                riepilogo[__fdn]["titolo"] = __ftit
                            riepilogo[__fdn]["real_code"] = True
                            riepilogo[__fdn]["NESSUN"] = True
                    for __v in riepilogo.values():
                        if not __v.get("real_code"):
                            continue
                        if (not __v.get("NC")
                                and not __v.get("OSS")
                                and not __v.get("NESSUN")):
                            __v["NESSUN"] = True

                else:
                    # CASO C: titolo generico (es. 'RILIEVO GENERALE') con NC/OSS
                    # -> mantieni comportamento storico: voce generica nel riepilogo.
                    __dn_p = normalize_code(__title_p)
                    if __dn_p not in riepilogo:
                        riepilogo[__dn_p] = {
                            "code": __title_p, "titolo": "", "rev": "",
                            "NC": False, "OSS": False, "NESSUN": False,
                            "real_code": False}
                    if "OSS" in __tag_n_p:
                        riepilogo[__dn_p]["OSS"] = True
                    else:
                        riepilogo[__dn_p]["NC"] = True
            # ─────────────────────────────────────────────────────────────────────

            # Mappa: indice riga tabella rilievi -> codice interno (NC-<Label>/OSS-<Label>)
            row_code_map = {}

            # immagini disciplina
            images_dir_disc = os.path.join(images_root, sanitize_filename(disciplina, "DISCIPLINA"))

            # mappa per appendice immagini
            code_to_images = {}
            code_to_desc = {}

            # Ordina: tutte le NC/OSS dello stesso elaborato consecutive
            def _label_num(x: str) -> int:
                m = re.search(r"(\d+)$", safe_str(x))
                return int(m.group(1)) if m else 999999

            # calcola tipo e label per ordinamento
            tmp_sort = []
            for __i, __r in subset.iterrows():
                __title = safe_str(__r.get(col_tit, ""))
                __tag = safe_str(__r.get(col_tag, ""))
                __tag_norm = normalize_text(__tag)
                __tipo_ord = 1 if "OSS" in __tag_norm else 0  # NC prima
                __label = safe_str(__r.get(col_label, "")) if col_label else ""
                if not __label:
                    m = re.search(r"\b[A-Z]{1,4}\d{1,4}[-_]\d{1,4}\b", normalize_text(__title))
                    __label = m.group(0).replace("_", "-") if m else ""
                tmp_sort.append((__title, __tipo_ord, _label_num(__label), __i))
            tmp_sort.sort(key=lambda t: (t[0], t[1], t[2], t[3]))
            subset_sorted = subset.loc[[t[3] for t in tmp_sort]]

            for _, r in subset_sorted.iterrows():
                title = safe_str(r.get(col_tit, ""))
                descr = safe_str(r.get(col_des, ""))
                tag = safe_str(r.get(col_tag, ""))
                status_raw = safe_str(r.get(col_sta, ""))


                # Titolo/Revisione elaborato (serve anche per i ToDo "NESSUN RILIEVO" che vanno nel riepilogo)
                titolo_doc, revisione = find_titolo_elaborato(title, report_lookup=lookup, elenco_lookup=elenco_lookup)

                # Recupera doc_norm dal pre-passo (reale o fallback su titolo).
                doc_code = extract_document_code_from_title(title) or title
                doc_norm = normalize_code(doc_code)

                if is_nessun_rilievo(tag, title):
                    # Nessun rilievo: non inserire riga in tabella rilievi.
                    # Il documento è già nel riepilogo dal pre-passo con NESSUN=True.
                    continue


                # Label (nuova regola): preferisci colonna Label
                label = safe_str(r.get(col_label, "")) if col_label else ""
                if not label:
                    m = re.search(r"\b[A-Z]{1,4}\d{1,4}[-_]\d{1,4}\b", normalize_text(title))
                    label = m.group(0).replace("_", "-") if m else sanitize_filename(title, "TODO")[:30]

                tag_norm = normalize_text(tag)
                if "OSS" in tag_norm:
                    tipo = "OSS"
                    oss_count += 1
                else:
                    tipo = "NC"
                    nc_count += 1

                codice_internal = f"{tipo}-{label}"
                seq_num = oss_count if tipo == "OSS" else nc_count
                codice_main = f"{tipo}{seq_num}"
                codice_sub = f"({tr_from_label(label)})"
                stato = traduci_stato(tag, status_raw)


                # immagini (salvate rinominate)
                image_paths = gather_todo_images(
                    r, label=label, images_dir=images_dir_disc,
                    image_cols=image_cols, photo_dir=photo_dir, disciplina=disciplina
                )

                # RIGA in tabella rilievi
                row = tab_rilievi.add_row().cells

                def set_cell(idx, val):
                    if idx < len(row):
                        row[idx].text = val

                # mapping colonne atteso (dal tuo screenshot):
                # 0: NC/OSS
                # 1: CODICE ELABORATO  -> Title ToDo (di solito il codice documento)
                # 2: TITOLO ELABORATO  -> da report
                # 3: RILIEVI ODI       -> descrizione ToDo
                # 4: ISPETTORE
                # ...
                write_code_cell(row[0], codice_main, codice_sub, main_pt=10, sub_pt=8)
                set_cell(1, title)
                set_cell(2, titolo_doc)
                set_cell(3, descr)
                # ispettore: usa SEMPRE il redattore da ELENCO ELABORATI (per disciplina/Codice_SP).
                # Se non coincide con "Created by" del ToDo, manteniamo comunque quanto indicato in ELENCO ELABORATI.
                # Fallback: se il redattore non è disponibile, usa il valore del ToDo.
                # Ispettore per singola NC/OSS:
                # PRIORITÀ = ToDo "Created by" (chi ha creato il rilievo, tipicamente l'ispettore).
                # Fallback = ToDo "Last modified by" (se manca Created by).
                # Ulteriori fallback (solo se ToDo vuoto) = mappa doc->isp (file gruppi) -> redattore disciplina.
                isp_created = safe_str(r.get("Created_by", r.get("Created by", r.get("created_by", "")))).strip()
                isp_lastmod = safe_str(
                    r.get("Last_modified_by",
                          r.get(col_lastmod, r.get("Last modified by", r.get("last_modified_by", ""))))
                ).strip()
                if isp_created:
                    isp = isp_created
                elif isp_lastmod:
                    isp = isp_lastmod
                else:
                    doc_code_for_isp = extract_document_code_from_title(title) or title
                    isp_from_map = find_ispettore_for_doc(doc_code_for_isp, doc_to_isp)
                    isp = isp_from_map or primary_redattore(nome_redattore_disc) or safe_str(r.get("Created by", ""))
                # Applica la mappa ispettori per disciplina
                isp = apply_ispettore_override(isp, disciplina)
                # v8.6.5: converti in iniziali SOLO per la cella della tabella rilievi
                isp_initials = to_inspector_initials(isp, initials_map, _ISPETTORE_ALIAS)
                set_cell(4, isp_initials)
                # Salva il codice interno per poter creare i link (se immagini) senza perdere la formattazione della cella
                row_code_map[len(tab_rilievi.rows)-1] = codice_internal
                # stato: spesso ultima colonna; qui mettiamo in 7 se esiste
                if len(row) > 7:
                    row[7].text = stato

                # Memorizza per appendice immagini (solo se immagini presenti)
                if image_paths:
                    code_to_images[codice_internal] = image_paths
                    # caption = codice + descrizione (pulita)
                    desc_clean = re.sub(r"\s+", " ", descr).strip()
                    if len(desc_clean) > 240:
                        desc_clean = desc_clean[:237] + "..."
                    code_to_desc[codice_internal] = desc_clean

                # (NC/OSS già aggiornati nel pre-passo per i codici reali; nessuna azione aggiuntiva)

            # riepilogo documenti: stampa TUTTI i documenti della disciplina.
            # Ordina per codice elaborato.
            assenza_count = 0
            for _, v in sorted(riepilogo.items(), key=lambda kv: safe_str(kv[1].get("code", ""))):
                # Nell'ULTIMA tabella devono comparire solo veri elaborati.
                # Voci generiche tipo "RILIEVO GENERALE" o pseudo-cartelle non
                # sono codici elaborato e restano visibili solo nelle righe della scheda.
                if not v.get("real_code"):
                    continue
                has_nc     = bool(v.get("NC"))
                has_oss    = bool(v.get("OSS"))
                # ASSENZA NC/OSS = X solo se il documento ha "Nessun rilievo"
                # e NON ha alcuna NC/OSS nello stesso riepilogo.
                has_nessun = bool(v.get("NESSUN"))
                assenza_nc_oss = has_nessun and not (has_nc or has_oss)
                if assenza_nc_oss:
                    assenza_count += 1
                clean_code = strip_revision_suffix(v.get("code", ""))
                rev_val = format_revision_value(v.get("rev", ""), v.get("code", ""))
                vals = [
                    clean_code,                            # 0: CODICE ELABORATO
                    rev_val,                               # 1: REVISIONE
                    safe_str(v.get("titolo", "")),        # 2: TITOLO ELABORATO
                    "X" if has_nc else "",                # 3: PRESENZA DI NC
                    "X" if has_oss else "",               # 4: PRESENZA DI OSS
                    "X" if assenza_nc_oss else "",        # 5: ASSENZA NC/OSS
                ]
                try:
                    cells = tab_riep.add_row().cells
                    for col_i, val in enumerate(vals):
                        if col_i < len(cells):
                            cells[col_i].text = val
                except Exception:
                    pass


            # Appendice immagini + link interni (merge "secondo script immagini")
            add_images_appendix_and_links(doc, tab_rilievi, code_to_images, code_to_desc, row_code_map)

            # footer
            # Conta solo i documenti con codice elaborato reale (esclude voci generiche
            # come "RILIEVI GENERALI" che non corrispondono a un elaborato specifico).
            docs_verificati = sum(1 for v in riepilogo.values() if v.get("real_code"))
            p = doc.add_paragraph()
            run = p.add_run(
                f"Riepilogo rilievi: NC = {nc_count}; OSS = {oss_count}; "
                f"Documenti verificati = {docs_verificati}")
            run.bold = True
            run.font.size = Pt(11)
            p.alignment = 1

            out_file = os.path.join(out_dir, f"{sanitize_filename(codice_sp_disc,'SP')}_{sanitize_filename(disciplina,'DISC')}.docx")
            def save_with_fallback_V2(doc, out_file: str):
                base, ext = os.path.splitext(out_file)
                try:
                    doc.save(out_file)
                    return out_file
                except PermissionError:
                    # File aperto o bloccato -> salva con suffisso _V2, _V3, ecc.
                    for i in range(2, 100):
                        alt_file = f"{base}_V{i}{ext}"
                        try:
                            doc.save(alt_file)
                            return alt_file
                        except PermissionError:
                            continue
                    raise

            saved_path = save_with_fallback_V2(doc, out_file)
            try:
                os.remove(tmp)
            except Exception:
                pass

        messagebox.showinfo("Completato", "Schede ispettive generate correttamente (titolo elaborato + immagini in appendice).")


# ============================================================================
# CLI / BACKEND MODE
# ============================================================================
class _CLIValue:
    """Sostituto minimo di tk.StringVar per usare SchedeGUI.emissione_schede
    senza creare finestre Tkinter.
    """
    def __init__(self, value=""):
        self._value = value or ""

    def get(self):
        return self._value

    def set(self, value):
        self._value = value or ""


def _find_first_file(folder: str, patterns: list[str], exts: tuple[str, ...]) -> str:
    folder = os.path.abspath(folder)
    if not os.path.isdir(folder):
        return ""
    files = []
    for fn in os.listdir(folder):
        low = fn.lower()
        if fn.startswith("~"):
            continue
        if not low.endswith(exts):
            continue
        files.append(os.path.join(folder, fn))
    for pat in patterns:
        pat_low = pat.lower()
        for p in files:
            if pat_low in os.path.basename(p).lower():
                return p
    return files[0] if files else ""


def _patch_messagebox_for_cli():
    """Evita popup Tkinter in esecuzione backend/CLI."""
    def _showinfo(title, message):
        print(f"[INFO] {title}: {message}")
        return True

    def _showwarning(title, message):
        print(f"[WARNING] {title}: {message}", file=sys.stderr)
        return True

    def _showerror(title, message):
        raise RuntimeError(f"{title}: {message}")

    messagebox.showinfo = _showinfo
    messagebox.showwarning = _showwarning
    messagebox.showerror = _showerror


def genera_schede_cli(
    elenco: str,
    todo: str,
    report: str,
    template: str,
    data_ricezione: str = "",
    data_emissione: str = "",
) -> str:
    """Genera le schede ispettive senza GUI.

    Ritorna la cartella output:
        <cartella_elenco>/Schede generate da ToDo
    """
    elenco = os.path.abspath(elenco or "")
    todo = os.path.abspath(todo or "")
    report = os.path.abspath(report or "")
    template = os.path.abspath(template or "")

    missing = []
    if not elenco or not os.path.isfile(elenco):
        missing.append("ELENCO ELABORATI (.xlsx)")
    if not todo or not os.path.isfile(todo):
        missing.append("ToDo Trimble export (.xlsx)")
    if not report or not os.path.isfile(report):
        missing.append("Report coerenze/elaborati (.xlsx)")
    if not template or not os.path.isfile(template):
        missing.append("Template Word scheda (.docx)")
    if missing:
        raise FileNotFoundError("File mancanti: " + ", ".join(missing))

    _patch_messagebox_for_cli()

    gui = SchedeGUI.__new__(SchedeGUI)
    gui.files = {
        "elenco": _CLIValue(elenco),
        "todo": _CLIValue(todo),
        "report": _CLIValue(report),
        "template": _CLIValue(template),
    }
    gui.data_ricezione_override = _CLIValue(data_ricezione)
    gui.data_emissione_override = _CLIValue(data_emissione)

    SchedeGUI.emissione_schede(gui)

    return os.path.join(os.path.dirname(elenco), "Schede generate da ToDo")


def _main_cli(argv=None) -> int:
    import argparse

    parser = argparse.ArgumentParser(
        description="Genera schede ispettive Word da ELENCO, ToDo Trimble, Report e Template."
    )
    parser.add_argument("--elenco", default="", help="Percorso ELENCO ELABORATI .xlsx")
    parser.add_argument("--todo", default="", help="Percorso ToDo Trimble export .xlsx")
    parser.add_argument("--report", default="", help="Percorso Report coerenze/elaborati .xlsx")
    parser.add_argument("--template", default="", help="Percorso Template Word .docx")
    parser.add_argument("--data-ricezione", default="", help="Override data ricezione elaborati")
    parser.add_argument("--data-emissione", default="", help="Override data emissione scheda")
    parser.add_argument(
        "--workdir",
        default=os.getcwd(),
        help="Cartella in cui cercare automaticamente i file se non indicati esplicitamente",
    )
    parser.add_argument(
        "--gui",
        action="store_true",
        help="Forza apertura GUI Tkinter",
    )

    args = parser.parse_args(argv)

    if args.gui:
        root = tk.Tk()
        SchedeGUI(root)
        root.mainloop()
        return 0

    workdir = os.path.abspath(args.workdir)

    elenco = args.elenco or _find_first_file(
        workdir,
        ["elenco", "elaborati"],
        (".xlsx",),
    )
    todo = args.todo or _find_first_file(
        workdir,
        ["todo", "trimble"],
        (".xlsx",),
    )
    report = args.report or _find_first_file(
        workdir,
        ["report", "coeren", "verifica"],
        (".xlsx",),
    )
    template = args.template or _find_first_file(
        workdir,
        ["template", "scheda"],
        (".docx",),
    )

    out_dir = genera_schede_cli(
        elenco=elenco,
        todo=todo,
        report=report,
        template=template,
        data_ricezione=args.data_ricezione,
        data_emissione=args.data_emissione,
    )

    print(f"[OK] Schede generate in: {out_dir}")
    try:
        for fn in sorted(os.listdir(out_dir)):
            if fn.lower().endswith(".docx"):
                print(os.path.join(out_dir, fn))
    except Exception:
        pass
    return 0


if __name__ == "__main__":
    # Compatibilita:
    # - senza argomenti, prova modalita CLI con auto-detect nella cartella corrente;
    # - se vuoi la vecchia GUI, lancia: python genera_schede_ispettive_v8_6_7_corretto.py --gui
    try:
        raise SystemExit(_main_cli())
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        raise SystemExit(1)
